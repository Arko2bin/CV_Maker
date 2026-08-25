from __future__ import annotations

import io, json, math, re, time
from copy import deepcopy
from typing import Any, Dict, List, Tuple
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader

RESUME_TEMPLATE: Dict[str, Any] = {"name":"","headline":"","email":"","phone":"","location":"","linkedin":"","summary":"","skills":[],"experience":[],"education":[],"projects":[],"certifications":[]}
SYSTEM_RULES = "Return only valid JSON. Never invent or infer unsupported employers, dates, degrees, certifications, technologies, duties, achievements, or metrics. Preserve the candidate's meaning."

class ProviderError(RuntimeError): pass

def resume_evidence_chunks(resume: Dict[str, Any]) -> List[Dict[str, str]]:
    """Create factual, section-level RAG documents from Qwen's structured extraction."""
    chunks: List[Dict[str, str]] = []
    for item in resume.get("experience", []):
        label = " | ".join(str(item.get(key, "")) for key in ("role", "company", "start_date", "end_date") if item.get(key))
        body = " ".join(str(bullet) for bullet in item.get("bullets", []))
        if label or body: chunks.append({"source": "experience: " + label, "text": label + ". " + body})
    for item in resume.get("projects", []):
        label = " | ".join(str(item.get(key, "")) for key in ("name", "technologies") if item.get(key))
        body = " ".join(str(bullet) for bullet in item.get("bullets", []))
        if label or body: chunks.append({"source": "project: " + label, "text": label + ". " + body})
    skills = ", ".join(str(skill) for skill in resume.get("skills", []))
    if skills: chunks.append({"source": "candidate skills", "text": skills})
    for item in resume.get("education", []):
        text = " | ".join(str(item.get(key, "")) for key in ("degree", "institution", "details") if item.get(key))
        if text: chunks.append({"source": "education", "text": text})
    return chunks

def create_embeddings(texts: List[str], api_key: str, model="text-embedding-3-small") -> List[List[float]]:
    if not texts: return []
    try:
        response = requests.post("https://api.openai.com/v1/embeddings", headers={"Authorization":"Bearer "+api_key,"Content-Type":"application/json"}, json={"model":model,"input":texts,"encoding_format":"float"}, timeout=120)
    except requests.RequestException as error: raise ProviderError("Could not create RAG embeddings: "+str(error))
    if not response.ok: raise ProviderError("Embedding provider HTTP {}: {}".format(response.status_code,response.text[:500]))
    try: return [item["embedding"] for item in sorted(response.json()["data"], key=lambda item:item["index"])]
    except (KeyError, TypeError, ValueError): raise ProviderError("Unexpected embedding response.")

def cosine_similarity(left: List[float], right: List[float]) -> float:
    denominator = math.sqrt(sum(value * value for value in left)) * math.sqrt(sum(value * value for value in right))
    return sum(a * b for a, b in zip(left, right)) / denominator if denominator else 0.0

class InMemoryVectorStore:
    """Per-request semantic vector store; it retains no resume data after the request."""
    def __init__(self): self.records: List[Dict[str, Any]] = []
    def add(self, chunks: List[Dict[str, str]], embeddings: List[List[float]]) -> None:
        if len(chunks) != len(embeddings): raise ProviderError("RAG index received mismatched chunks and embeddings.")
        self.records = [{"source":chunk["source"],"evidence":chunk["text"],"embedding":embedding} for chunk,embedding in zip(chunks,embeddings)]
    def query(self, query_embeddings: List[List[float]], limit=8) -> List[Dict[str, Any]]:
        scored: Dict[int, float] = {}
        for index, record in enumerate(self.records):
            scored[index] = max([cosine_similarity(record["embedding"], query) for query in query_embeddings] or [0.0])
        ranked = sorted(scored.items(), key=lambda pair:pair[1], reverse=True)[:limit]
        return [{"source":self.records[index]["source"],"evidence":self.records[index]["evidence"],"similarity":round(score,4)} for index,score in ranked if score > 0]

def build_rag_context(preprocessed: Dict[str, Any], api_key: str) -> Dict[str, Any]:
    """Chunk, embed, store, and semantically retrieve factual evidence for current JD requirements."""
    jd = preprocessed.get("jd_analysis", {})
    queries = []
    for key in ("required_skills", "preferred_skills", "responsibilities", "keywords", "core_capabilities"):
        value = jd.get(key, [])
        queries.extend(value if isinstance(value, list) else [str(value)])
    for item in preprocessed.get("capability_evidence", []):
        if isinstance(item, dict): queries.append(str(item.get("capability", "")))
    chunks = resume_evidence_chunks(preprocessed["resume"])
    queries = [query for query in queries[:60] if query.strip()]
    store = InMemoryVectorStore()
    store.add(chunks, create_embeddings([chunk["text"] for chunk in chunks], api_key))
    return {"retrieval_queries":queries, "index_size":len(store.records), "retrieved_candidate_evidence":store.query(create_embeddings(queries, api_key))}

def extract_document(file) -> str:
    data, ext = file.getvalue(), file.name.lower().rsplit(".",1)[-1]
    if ext == "docx":
        doc = Document(io.BytesIO(data)); parts = [p.text for p in doc.paragraphs]
        for table in doc.tables: parts += [" | ".join(c.text for c in row.cells) for row in table.rows]
        text = "\n".join(x for x in parts if x.strip())
    elif ext == "pdf": text = "\n".join(p.extract_text() or "" for p in PdfReader(io.BytesIO(data)).pages)
    elif ext == "txt": text = data.decode("utf-8", errors="replace")
    else: raise ValueError("Upload a DOCX, PDF, or TXT file.")
    if not text.strip(): raise ValueError("No readable text found. Image-only PDFs need OCR.")
    return text.strip()

def parse_json(text: str) -> Dict[str, Any]:
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", text.strip(), flags=re.I)
    decoder = json.JSONDecoder()
    for i, char in enumerate(cleaned):
        if char == "{":
            try:
                value, _ = decoder.raw_decode(cleaned[i:])
                if isinstance(value, dict): return value
            except json.JSONDecodeError: pass
    raise ProviderError("The model returned invalid JSON. Try again.")

def normalize(data: Dict[str, Any]) -> Dict[str, Any]:
    result = deepcopy(RESUME_TEMPLATE)
    for key in result:
        if data.get(key) is not None: result[key] = data[key]
    for key in ("skills","experience","education","projects","certifications"):
        if not isinstance(result[key], list): result[key] = []
    return result

def post_chat(url: str, token: str, model: str, messages: List[Dict[str,str]], json_mode=False) -> str:
    payload: Dict[str,Any] = {"model":model,"messages":messages,"temperature":0.15}
    if json_mode: payload["response_format"] = {"type":"json_object"}
    response = None
    # Hosted inference providers occasionally time out before the model starts.
    # Retry only transient gateway failures, with a short bounded backoff.
    for attempt in range(3):
        try:
            response = requests.post(url, headers={"Authorization":"Bearer "+token,"Content-Type":"application/json"}, json=payload, timeout=240)
        except requests.RequestException as error:
            if attempt == 2: raise ProviderError("Could not reach provider after 3 attempts: "+str(error))
            time.sleep(2 ** attempt)
            continue
        if response.status_code not in (502, 503, 504) or attempt == 2:
            break
        time.sleep(2 ** attempt)
    if response is None:
        raise ProviderError("Could not reach provider.")
    if not response.ok:
        try: detail = response.json().get("error", response.json())
        except ValueError: detail = response.text[:500]
        if response.status_code == 403 and "huggingface.co" in url:
            raise ProviderError(
                "Hugging Face rejected this token. Create a fine-grained token and enable "
                "'Make calls to Inference Providers', then replace HF_TOKEN in "
                ".streamlit/secrets.toml and restart Streamlit."
            )
        if response.status_code in (502, 503, 504) and "huggingface.co" in url:
            raise ProviderError(
                "The resume-analysis service is temporarily busy or timed out after 3 attempts. "
                "Please try again in a few minutes, or use a shorter resume/job description."
            )
        raise ProviderError("Provider HTTP {}: {}".format(response.status_code, detail))
    try: return response.json()["choices"][0]["message"]["content"]
    except (KeyError,IndexError,TypeError,ValueError): raise ProviderError("Unexpected provider response.")

def preprocess_qwen(source: str, job: str, token: str, model: str) -> Dict[str,Any]:
    prompt = """Perform factual extraction and job-description analysis for a resume-tailoring system.

Return one JSON object with these keys:
- resume: the resume represented by the exact RESUME SHAPE below
- jd_analysis: an object containing target_job_title, target_professional_identity, required_skills[], preferred_skills[], responsibilities[], keywords[], qualifications[], and core_capabilities[]
- skill_categories: an object grouping only the candidate's explicitly proven skills into technical[], tools[], domain[], and soft_skills[]
- capability_evidence: an array of objects with capability, supporting_resume_evidence, and demonstrated_technologies; capture transferable capabilities such as system design, APIs, automation, backend development, data processing, architecture, debugging, testing, and integration only when supported

Extract facts and analyze the JD only. Do not tailor, reframe, rewrite, improve, summarize, infer, or add candidate information.
Experience objects: company, role, location, start_date, end_date, bullets[].
Education objects: institution, degree, location, start_date, end_date, details[].
Project objects: name, technologies, link, bullets[]. Certifications: name, issuer, date.

RESUME SHAPE: {}

SPECIFIC RESUME DATA:
{}

TARGET JOB DESCRIPTION:
{}""".format(json.dumps(RESUME_TEMPLATE), source[:35000], job[:25000])
    answer = post_chat("https://router.huggingface.co/v1/chat/completions",token,model,[{"role":"system","content":SYSTEM_RULES},{"role":"user","content":prompt}])
    result = parse_json(answer)
    result["resume"] = normalize(result.get("resume", {}))
    return result

def finalize_openai(preprocessed: Dict[str,Any], key: str, model: str) -> Dict[str,Any]:
    rag_context = build_rag_context(preprocessed,key)
    prompt = """You are the final expert resume architect. Build the resume from the TARGET JD, not from the source resume's existing labels or section wording. Treat the extracted candidate data only as factual evidence of capabilities, not as a template to preserve.

The target JD determines the resume's target job title, professional positioning, summary, skills emphasis, bullet structure, ordering, and terminology. Reconstruct every section around the target role rather than editing the extracted resume section-by-section.

TRANSFERABLE-EXPERIENCE RULE:
- Translate proven underlying capabilities into concepts relevant to the target role aggressively but truthfully.
- For example, supported Python/AI work may evidence system design, APIs, automation, backend development, data processing, architecture, debugging, testing, or integration. Emphasize those underlying engineering capabilities for a Java/backend JD.
- Apply the same evidence-based mapping for transitions such as PHP to Ruby on Rails or Python to Java.
- Do not mechanically preserve the source technology, headline, summary, skill order, or bullet wording.
- Center the professional identity on the target role. However, a target-facing headline must not state or imply hands-on experience with a specific technology that is absent from the evidence. Prefer a truthful bridge such as "Backend Engineer | Targeting Java Ecosystems" over falsely claiming "Experienced Java Engineer".
- Clearly distinguish transferable capability from demonstrated technology experience. Never list an unsupported target technology as a candidate skill or claim it was used in employment or projects.

Do not copy the extracted resume or make only minor wording changes. You must transform its presentation while preserving factual truth:
1. SELECT only the experience, projects, skills, and evidence that best support the target role.
2. PRIORITIZE the strongest JD matches by reordering skills, experience bullets, and projects.
3. REFRAME supported experience around the JD's responsibilities and desired outcomes without changing what the candidate actually did.
4. REWRITE bullets substantially using concise action-led professional language. Change sentence structure, emphasis, and organization—not merely synonyms.
5. CONSOLIDATE overlapping or repetitive bullets and omit low-value irrelevant bullets. Do not preserve the original bullet count.
6. Write a new targeted 3-4 sentence summary that connects the candidate's strongest proven evidence to the target role.
7. Use JD keywords naturally only when candidate evidence explicitly supports them. Absence from candidate evidence means the claim is prohibited.
8. Preserve identity, contact details, employers, job titles, dates, degrees, and certifications exactly as extracted.
9. Never invent or imply technologies, duties, achievements, metrics, employers, dates, qualifications, or experience.
10. Run a final grammar, clarity, repetition, and ATS-readability check before returning the result.
11. Reconstruct the skills section around JD capability groups, but include only technologies and skills supported by candidate evidence.
12. Ensure every rewritten bullet remains traceable to at least one item in capability_evidence or the extracted resume.
13. Use RETRIEVED CANDIDATE EVIDENCE as the primary evidence set for job-specific bullets. If no retrieved evidence supports a JD keyword, omit that keyword rather than making a claim.

The final result must look purpose-built for the target JD, with the extracted CV serving only as factual evidence.

Return the complete final resume as valid JSON in exactly the RESUME SHAPE shown below. Return JSON only.

RESUME SHAPE:
{}

QWEN FACTUAL EXTRACTION AND JD ANALYSIS (the only permitted source of candidate facts):
{}

RAG RETRIEVED CANDIDATE EVIDENCE (selected locally from the structured extraction for the current JD):
{}""".format(json.dumps(RESUME_TEMPLATE),json.dumps(preprocessed,ensure_ascii=False),json.dumps(rag_context,ensure_ascii=False))
    answer = post_chat("https://api.openai.com/v1/chat/completions",key,model,[{"role":"system","content":SYSTEM_RULES},{"role":"user","content":prompt}],True)
    return normalize(parse_json(answer))

def run_pipeline(source: str, job: str, openai_key: str, hf_token: str, openai_model: str, hf_model: str) -> Tuple[Dict[str,Any],Dict[str,Any]]:
    preprocessed = preprocess_qwen(source,job,hf_token,hf_model)
    return preprocessed["resume"], finalize_openai(preprocessed,openai_key,openai_model)

def ats_analysis(resume: Dict[str,Any], job: str) -> Tuple[int,List[str],List[str]]:
    stop={"and","the","with","for","that","this","from","your","you","our","are","will","have","has","job","role","work","team","years","using","skills","experience","required","preferred"}
    counts: Dict[str,int]={}
    for word in re.findall(r"[a-zA-Z][a-zA-Z+#.\-]{2,}",job.lower()):
        if word not in stop: counts[word]=counts.get(word,0)+1
    terms=[x[0] for x in sorted(counts.items(),key=lambda x:(-x[1],x[0]))[:30]]; content=json.dumps(resume).lower()
    matched=[x for x in terms if re.search(r"(?<!\w){}(?!\w)".format(re.escape(x)),content)]; missing=[x for x in terms if x not in matched]
    complete=sum(bool(resume.get(k)) for k in ("name","email","summary","skills","experience","education"))
    return min(100,round(70*len(matched)/max(1,len(terms)))+5*complete),matched,missing

def heading(doc: Document, title: str):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(9); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(title.upper()); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(29,78,121)
    borders=OxmlElement("w:pBdr"); bottom=OxmlElement("w:bottom")
    for k,v in (("val","single"),("sz","6"),("space","1"),("color","1D4E79")): bottom.set(qn("w:"+k),v)
    borders.append(bottom); p._p.get_or_add_pPr().append(borders)

def bullets(doc: Document, items):
    for item in items or []: doc.add_paragraph(str(item),style="List Bullet").paragraph_format.space_after=Pt(1)

def build_docx(resume: Dict[str,Any]) -> bytes:
    doc=Document(); sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=Inches(.55); sec.left_margin=sec.right_margin=Inches(.7)
    doc.styles["Normal"].font.name="Arial"; doc.styles["Normal"].font.size=Pt(9.5)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(str(resume.get("name") or "Your Name")); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor(29,78,121)
    for value in [resume.get("headline"),"  •  ".join(str(resume.get(k)) for k in ("email","phone","location","linkedin") if resume.get(k))]:
        if value: p=doc.add_paragraph(str(value)); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if resume.get("summary"): heading(doc,"Professional Summary"); doc.add_paragraph(str(resume["summary"]))
    if resume.get("skills"): heading(doc,"Core Skills"); doc.add_paragraph(" • ".join(map(str,resume["skills"])))
    for section,title,bkey,namekeys,extra_keys in [
        ("experience","Professional Experience","bullets",("role","company"),("location","start_date","end_date")),
        ("projects","Projects","bullets",("name","technologies"),("link",)),
        ("education","Education","details",("degree","institution"),("location","start_date","end_date")),
    ]:
        if resume.get(section):
            heading(doc,title)
            for item in resume[section]:
                p=doc.add_paragraph(); p.add_run(" — ".join(str(item.get(k,"")) for k in namekeys if item.get(k))).bold=True
                extra=" | ".join(str(item.get(k,"")) for k in extra_keys if item.get(k))
                if extra: p.add_run(" | "+extra)
                bullets(doc,item.get(bkey,[]))
    if resume.get("certifications"):
        heading(doc,"Certifications")
        for item in resume["certifications"]: doc.add_paragraph(" — ".join(str(item.get(k,"")) for k in ("name","issuer","date") if item.get(k)))
    output=io.BytesIO(); doc.save(output); return output.getvalue()
